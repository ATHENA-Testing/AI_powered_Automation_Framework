"""
File Processor for Python Automation Framework
Handles extraction of text content from various file formats.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
import tempfile

# Document processing imports
import PyPDF2
from docx import Document
import openpyxl
import pandas as pd

from src.config_manager import ConfigManager
from src.logger import log, LoggerMixin

class FileProcessor(LoggerMixin):
    """
    Processes various file formats and extracts text content.
    """

    def __init__(self, config: ConfigManager):
        """
        Initializes the FileProcessor.

        Args:
            config: An instance of ConfigManager for accessing configuration settings.
        """
        self.config = config
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.txt': self._extract_from_txt,
            '.doc': self._extract_from_doc,
            '.docx': self._extract_from_docx,
            '.xlsx': self._extract_from_xlsx,
            '.xls': self._extract_from_xls,
            '.csv': self._extract_from_csv
        }
        
        self.logger.info("FileProcessor initialized")

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extracts text content from a file based on its extension.

        Args:
            file_path: Path to the file to process.

        Returns:
            Extracted text content as a string.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        self.logger.info(f"Extracting text from {file_path} (format: {file_extension})")
        
        try:
            extractor_function = self.supported_formats[file_extension]
            content = extractor_function(str(file_path))
            
            self.logger.info(f"Successfully extracted {len(content)} characters from {file_path}")
            return content
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extracts text from a PDF file.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Extracted text content.
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
            
            return "\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise

    def _extract_from_txt(self, file_path: str) -> str:
        """
        Extracts text from a plain text file.

        Args:
            file_path: Path to the text file.

        Returns:
            File content as string.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                self.logger.error(f"Error reading text file {file_path}: {e}")
                raise
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_from_doc(self, file_path: str) -> str:
        """
        Extracts text from a DOC file.
        Note: This is a placeholder. Full DOC support would require additional libraries.

        Args:
            file_path: Path to the DOC file.

        Returns:
            Extracted text content.
        """
        # For DOC files, we would need python-docx2txt or similar
        # For now, we'll return a placeholder message
        self.logger.warning(f"DOC file processing not fully implemented for {file_path}")
        return f"[DOC file content from {file_path} - processing not fully implemented]"

    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extracts text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Extracted text content.
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    def _extract_from_xlsx(self, file_path: str) -> str:
        """
        Extracts text from an XLSX file.

        Args:
            file_path: Path to the XLSX file.

        Returns:
            Extracted text content.
        """
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    
                    if row_text:  # Only add non-empty rows
                        text_content.append(" | ".join(row_text))
                
                text_content.append("")  # Add blank line between sheets
            
            return "\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from XLSX {file_path}: {e}")
            raise

    def _extract_from_xls(self, file_path: str) -> str:
        """
        Extracts text from an XLS file.

        Args:
            file_path: Path to the XLS file.

        Returns:
            Extracted text content.
        """
        try:
            # Use pandas to read XLS files
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content.append(f"Sheet: {sheet_name}")
                
                # Convert DataFrame to string representation
                text_content.append(df.to_string(index=False))
                text_content.append("")  # Add blank line between sheets
            
            return "\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from XLS {file_path}: {e}")
            raise

    def _extract_from_csv(self, file_path: str) -> str:
        """
        Extracts text from a CSV file.

        Args:
            file_path: Path to the CSV file.

        Returns:
            Extracted text content.
        """
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from CSV {file_path}: {e}")
            raise

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Gets information about a file.

        Args:
            file_path: Path to the file.

        Returns:
            Dictionary containing file information.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = file_path.stat()
        
        return {
            "name": file_path.name,
            "extension": file_path.suffix.lower(),
            "size_bytes": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "is_supported": file_path.suffix.lower() in self.supported_formats
        }

    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validates a file for processing.

        Args:
            file_path: Path to the file to validate.

        Returns:
            Dictionary containing validation results.
        """
        try:
            file_info = self.get_file_info(file_path)
            
            validation_result = {
                "is_valid": True,
                "file_info": file_info,
                "errors": [],
                "warnings": []
            }
            
            # Check if file format is supported
            if not file_info["is_supported"]:
                validation_result["is_valid"] = False
                validation_result["errors"].append(f"Unsupported file format: {file_info['extension']}")
            
            # Check file size (warn if > 50MB)
            if file_info["size_mb"] > 50:
                validation_result["warnings"].append(f"Large file size: {file_info['size_mb']} MB")
            
            # Check if file is empty
            if file_info["size_bytes"] == 0:
                validation_result["is_valid"] = False
                validation_result["errors"].append("File is empty")
            
            return validation_result
            
        except Exception as e:
            return {
                "is_valid": False,
                "file_info": None,
                "errors": [str(e)],
                "warnings": []
            }

    def batch_process_files(self, file_paths: list) -> Dict[str, Any]:
        """
        Processes multiple files in batch.

        Args:
            file_paths: List of file paths to process.

        Returns:
            Dictionary containing batch processing results.
        """
        results = {
            "successful": {},
            "failed": {},
            "total_files": len(file_paths),
            "success_count": 0,
            "failure_count": 0
        }
        
        for file_path in file_paths:
            try:
                content = self.extract_text_from_file(file_path)
                results["successful"][file_path] = {
                    "content": content,
                    "length": len(content)
                }
                results["success_count"] += 1
                
            except Exception as e:
                results["failed"][file_path] = str(e)
                results["failure_count"] += 1
        
        self.logger.info(f"Batch processing completed: {results['success_count']} successful, {results['failure_count']} failed")
        
        return results

# Example Usage (for testing purposes)
if __name__ == "__main__":
    # Assuming config.ini is in the parent directory
    current_dir = Path(__file__).parent
    config_path = current_dir.parent / "config.ini"
    
    # Initialize ConfigManager with the correct path
    cfg_manager = ConfigManager(config_path=str(config_path))
    file_processor = FileProcessor(cfg_manager)

    # Create a sample text file for testing
    test_file_path = current_dir.parent / "test_document.txt"
    with open(test_file_path, 'w') as f:
        f.write("This is a sample document for testing the file processor.\nIt contains multiple lines of text.\nThis will be used to test text extraction functionality.")

    print("\n--- File Processing Demo ---")
    
    # Test file validation
    print("1. Validating file...")
    validation_result = file_processor.validate_file(str(test_file_path))
    print(f"Validation result: {validation_result}")
    
    # Test text extraction
    print("\n2. Extracting text...")
    try:
        extracted_text = file_processor.extract_text_from_file(str(test_file_path))
        print(f"Extracted text ({len(extracted_text)} characters):")
        print(extracted_text)
    except Exception as e:
        print(f"Error extracting text: {e}")
    
    # Test file info
    print("\n3. Getting file info...")
    file_info = file_processor.get_file_info(str(test_file_path))
    print(f"File info: {file_info}")
    
    # Clean up test file
    test_file_path.unlink()
    
    print("\n--- File Processing Demo Complete ---")

